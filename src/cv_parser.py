"""Helpers for extracting useful CV information from uploaded files."""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import re
import subprocess
import sys
import tempfile

SKILL_ALIASES: dict[str, tuple[str, ...]] = {
    "Python": ("python",),
    "SQL": ("sql", "mysql", "sqlite", "t-sql", "tsql"),
    "React": ("react", "reactjs", "react.js"),
    "Django": ("django",),
    "FastAPI": ("fastapi", "fast api"),
    "PostgreSQL": ("postgresql", "postgres", "postgre sql"),
    "AWS": ("aws", "amazon web services"),
    "Docker": ("docker", "dockerized", "containers"),
    "JavaScript": ("javascript", "js"),
    "Excel": ("excel", "spreadsheets", "microsoft excel"),
    "TypeScript": ("typescript", "ts"),
    "Node.js": ("node.js", "nodejs", "node js"),
}

def _build_skill_patterns(
    aliases_map: dict[str, tuple[str, ...]],
) -> dict[str, tuple[re.Pattern[str], ...]]:
    """Build compiled regex patterns for known skill aliases."""
    result: dict[str, tuple[re.Pattern[str], ...]] = {}
    for canonical, aliases in aliases_map.items():
        patterns: list[re.Pattern[str]] = []
        for alias in aliases:
            escaped = re.escape(alias).replace(r"\ ", r"\s+")
            patterns.append(
                re.compile(
                    rf"(?<![A-Za-z0-9]){escaped}(?![A-Za-z0-9])",
                    flags=re.IGNORECASE,
                )
            )
        result[canonical] = tuple(patterns)
    return result


SKILL_PATTERNS = _build_skill_patterns(SKILL_ALIASES)


@dataclass(slots=True)
class CVAnalysis:
    """Structured summary extracted from a CV."""

    text: str = ""
    name: str = ""
    email: str = ""
    phone: str = ""
    linkedin_url: str = ""
    portfolio_url: str = ""
    skills: list[str] = field(default_factory=list)
    experience_summary: str = ""


@dataclass(slots=True)
class SkillMatchResult:
    """Skill overlap between CV content and one job."""

    matched: list[str] = field(default_factory=list)
    missing: list[str] = field(default_factory=list)
    inferred_job_skills: list[str] = field(default_factory=list)


def _fallback_python_candidates() -> list[Path]:
    """Return candidate Python interpreters that may have document parser deps."""
    home = Path.home()
    candidates = [
        Path(sys.executable),
        home / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python" / "python.exe",
        home / "AppData" / "Local" / "Python" / "pythoncore-3.14-64" / "python.exe",
    ]

    unique_candidates: list[Path] = []
    seen: set[str] = set()
    for candidate in candidates:
        normalized = str(candidate).lower()
        if normalized not in seen and candidate.exists():
            unique_candidates.append(candidate)
            seen.add(normalized)
    return unique_candidates


def _extract_with_external_python(data: bytes, mode: str) -> str:
    """Try extracting document text using another Python interpreter."""
    script = """
from io import BytesIO
from pathlib import Path
import sys

mode = sys.argv[1]
file_path = Path(sys.argv[2])
data = file_path.read_bytes()

if mode == "pdf":
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages]
    print("\\n".join(parts).strip())
elif mode == "docx":
    from docx import Document
    document = Document(BytesIO(data))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    print("\\n".join(parts).strip())
else:
    raise ValueError(f"Unsupported mode: {mode}")
""".strip()

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{mode}") as temp_file:
        temp_file.write(data)
        temp_path = Path(temp_file.name)

    try:
        for python_path in _fallback_python_candidates():
            try:
                result = subprocess.run(
                    [str(python_path), "-c", script, mode, str(temp_path)],
                    check=True,
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                )
                return result.stdout.strip()
            except subprocess.CalledProcessError:
                continue
    finally:
        temp_path.unlink(missing_ok=True)

    raise ValueError(
        f"{mode.upper()} CV support is not available in the current environment yet. "
        "Please install the required parser dependency in the Python environment used by Streamlit."
    )


def extract_text_from_pdf_bytes(data: bytes) -> str:
    """Extract plain text from a PDF byte stream."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return _extract_with_external_python(data, "pdf")

    reader = PdfReader(BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts).strip()


def extract_text_from_docx_bytes(data: bytes) -> str:
    """Extract plain text from a DOCX byte stream."""
    try:
        from docx import Document
    except ImportError:
        return _extract_with_external_python(data, "docx")

    document = Document(BytesIO(data))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(parts).strip()


def extract_known_skills(text: str) -> list[str]:
    """Extract canonical skill names from free text using alias-aware matching."""
    cleaned = text or ""
    found: list[str] = []
    for canonical, patterns in SKILL_PATTERNS.items():
        if any(pattern.search(cleaned) for pattern in patterns):
            found.append(canonical)
    return found


def _extract_name(lines: list[str]) -> str:
    """Guess a candidate name from the top of the CV."""
    banned_tokens = ("curriculum vitae", "resume", "cv", "profile", "summary")
    for line in lines[:5]:
        lowered = line.casefold()
        if any(token in lowered for token in banned_tokens):
            continue
        if "@" in line or "http" in lowered or re.search(r"\d", line):
            continue
        if 2 <= len(line.split()) <= 5:
            return line
    return ""


def _extract_experience_summary(cleaned: str, lines: list[str]) -> str:
    """Extract a short candidate summary from CV text."""
    summary_patterns = (
        r"(?:summary|profile|professional summary)\s*[:\-]\s*(.+)",
        r"(experienced .+?)(?:\.|\n|$)",
        r"(backend engineer.+?)(?:\.|\n|$)",
        r"(software engineer.+?)(?:\.|\n|$)",
        r"(data analyst.+?)(?:\.|\n|$)",
    )
    for pattern in summary_patterns:
        match = re.search(pattern, cleaned, flags=re.IGNORECASE | re.DOTALL)
        if match:
            sentence = re.sub(r"\s+", " ", match.group(1)).strip(" .")
            if sentence:
                return sentence

    for line in lines[1:8]:
        lowered = line.casefold()
        if "@" in line or "http" in lowered:
            continue
        if 6 <= len(line.split()) <= 20:
            return line.strip(" .")
    return ""


def analyze_cv_text(text: str) -> CVAnalysis:
    """Extract contact info and skills from plain CV text."""
    cleaned = re.sub(r"\r\n?", "\n", text or "").strip()
    lines = [line.strip() for line in cleaned.splitlines() if line.strip()]

    email_match = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", cleaned, flags=re.IGNORECASE)
    phone_match = re.search(r"(\+?\d[\d\s\-\(\)]{7,}\d)", cleaned)
    linkedin_match = re.search(r"https?://(?:www\.)?linkedin\.com/\S+", cleaned, flags=re.IGNORECASE)
    github_match = re.search(r"https?://(?:www\.)?github\.com/\S+", cleaned, flags=re.IGNORECASE)
    url_matches = re.findall(r"https?://\S+", cleaned, flags=re.IGNORECASE)

    portfolio_url = ""
    excluded_urls = {
        linkedin_match.group(0).rstrip(").,]") if linkedin_match else "",
        github_match.group(0).rstrip(").,]") if github_match else "",
    }
    for url in url_matches:
        normalized_url = url.rstrip(").,]")
        if normalized_url in excluded_urls:
            continue
        portfolio_url = normalized_url
        break

    skills = extract_known_skills(cleaned)

    return CVAnalysis(
        text=cleaned,
        name=_extract_name(lines),
        email=email_match.group(0) if email_match else "",
        phone=phone_match.group(1).strip() if phone_match else "",
        linkedin_url=linkedin_match.group(0).rstrip(").,]") if linkedin_match else "",
        portfolio_url=portfolio_url or (github_match.group(0).rstrip(").,]") if github_match else ""),
        skills=skills,
        experience_summary=_extract_experience_summary(cleaned, lines),
    )


def analyze_cv_bytes(file_name: str, data: bytes) -> CVAnalysis:
    """Analyze a PDF or DOCX CV from raw bytes."""
    suffix = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""
    if suffix == "pdf":
        text = extract_text_from_pdf_bytes(data)
    elif suffix == "docx":
        text = extract_text_from_docx_bytes(data)
    else:
        raise ValueError("Unsupported CV format. Please upload a PDF or DOCX file.")

    return analyze_cv_text(text)


def match_cv_skills_to_job(cv_skills: list[str], job_skills_text: str, job_description: str) -> SkillMatchResult:
    """Match CV skills against explicit and inferred job skills."""
    explicit_job_skills = [skill.strip() for skill in str(job_skills_text or "").split(";") if skill.strip()]
    inferred_job_skills = extract_known_skills(f"{job_skills_text} {job_description}")

    job_skill_set: list[str] = []
    for skill in explicit_job_skills + inferred_job_skills:
        canonical = next(
            (name for name in SKILL_ALIASES if name.casefold() == skill.casefold()),
            skill,
        )
        if canonical not in job_skill_set:
            job_skill_set.append(canonical)

    matched = [skill for skill in cv_skills if skill in job_skill_set]
    missing = [skill for skill in job_skill_set if skill not in cv_skills]
    return SkillMatchResult(
        matched=matched,
        missing=missing,
        inferred_job_skills=job_skill_set,
    )


def analyze_profile_text(text: str) -> CVAnalysis:
    """Analyze user-provided profile text (markdown or plain text).

    Supports formats like:

    Name: Santi Husni
    Email: santi@example.com
    Phone: +62812345678
    LinkedIn: https://linkedin.com/in/santi
    Skills: Python, FastAPI, PostgreSQL, Docker
    Experience: Backend engineer with 5 years...

    Also supports markdown headings.
    """
    # Strip markdown formatting for cleaner parsing
    cleaned = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)  # headings
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)        # bold
    cleaned = re.sub(r"\*(.*?)\*", r"\1", cleaned)             # italic
    cleaned = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", cleaned)      # links

    # Try labeled extraction first
    name = ""
    email = ""
    phone = ""
    linkedin = ""
    portfolio = ""
    raw_skills = ""
    experience_summary = ""

    patterns = {
        "name": r"(?i)(?:name|nama)[\s:]*\s*(.+)",
        "email": r"(?i)(?:email|e-mail|mail)[\s:]*\s*([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})",
        "phone": r"(?i)(?:phone|tel|telepon|no\.?\s*hp)[\s:]*\s*(\+?\d[\d\s\-()]{7,}\d)",
        "linkedin": r"(?i)(?:linkedin|linked\s*in)[\s:]*\s*(https?://(?:www\.)?linkedin\.com/\S+)",
        "portfolio": r"(?i)(?:portfolio|website|github)[\s:]*\s*(https?://\S+)",
        "skills": r"(?i)(?:skills?|keahlian|tech\s*stack)[\s:]*\s*(.+)",
        "experience": r"(?i)(?:experience|pengalaman|summary|about)[\s:]*\s*(.+)",
    }

    for field, pattern in patterns.items():
        match = re.search(pattern, cleaned, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            if field == "name":
                name = value
            elif field == "email":
                email = value
            elif field == "phone":
                phone = value
            elif field == "linkedin":
                linkedin = value
            elif field == "portfolio":
                portfolio = value
            elif field == "skills":
                raw_skills = value
            elif field == "experience":
                experience_summary = value

    # Fallback: use the main cv_text analyzer for unstructured text
    analysis = analyze_cv_text(text)

    # Merge: prefer labeled extraction, fallback to CV analyzer
    return CVAnalysis(
        text=text,
        name=name or analysis.name,
        email=email or analysis.email,
        phone=phone or analysis.phone,
        linkedin_url=linkedin or analysis.linkedin_url,
        portfolio_url=portfolio or analysis.portfolio_url,
        skills=extract_known_skills(f"{raw_skills} {text}")
            if raw_skills
            else analysis.skills,
        experience_summary=experience_summary or analysis.experience_summary,
    )
